
# importing modules
import urllib.request 
from bs4 import BeautifulSoup
from docx import Document
  
# providing url
url = "https://novelsonline.net/youkoso-jitsuryoku-shijou-shugi-no-kyoushitsu-e/volume-1/chapter-1-1-int"
  
# opening the url for reading
html = urllib.request.urlopen(url)
  
# parsing the html file
htmlParse = BeautifulSoup(html, 'html.parser')
  
# getting all the paragraphs
filename = 'newdoc.docx'
document = Document()
table = htmlParse.find("div", {"id": "contentall"})
for para in table.find_all("p"):
    if(para.string != None):
        print(para.string)
        document.add_paragraph(para.string)
#     document.add_paragraph(para.text)

document.save(filename)
